#!/usr/bin/env python3
"""
Генератор уведомлений о расчёте премии.

Использование:
    python generate.py путь_к_excel.xlsx

В Excel на листе «Расчет премии» укажите «Сделать» в столбце «Уведомление»
для строк, по которым нужно создать документ.

Файл template.docx должен лежать рядом со скриптом.
"""

VERSION = "2026-05-26d"   # tblW pct=5000 (AutoFit to Window)

import sys
import re
import os
from datetime import date, datetime
from pathlib import Path
import warnings

warnings.filterwarnings("ignore", category=UserWarning)

try:
    import openpyxl
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")

try:
    from num2words import num2words
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")


# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------

SCRIPT_DIR    = Path(__file__).parent
TEMPLATE_PATH = SCRIPT_DIR / "template.docx"


# ---------------------------------------------------------------------------
# Индексы столбцов по умолчанию (0-based).
# Используются как запасные значения, если заголовки в строке 2
# не распознаны автоматически.
# ---------------------------------------------------------------------------

_COL_DEFAULTS = {
    "date_from":  2,   # C — начало периода
    "date_to":    3,   # D — конец периода
    "client":     5,   # F — клиент
    "ds":         6,   # G — основание (ДС №... от ...)
    "comment":    7,   # H — комментарий (ключ легенды)
    "category":   8,   # I — категория (для маркетинг-фразы)
    "notif_num": 15,   # P — номер уведомления
    "turnover":  17,   # R — товарооборот (I)
    "exclusion": 18,   # S — вывод из-под бонуса (II)
    "returns":   19,   # T — обратная реализация (III)
    "overdue":   20,   # U — просроченная задолженность (IV)
    "samples":   21,   # V — отгруженные образцы (VII)
    "base":      22,   # W — база для начисления (V)
    "rate":      23,   # X — процент премии (VI)
    "bonus":     24,   # Y — сумма премии (VIII)
    "trigger":   25,   # Z — отметка «Сделать»
}

# Ключевые слова для поиска столбцов по заголовкам строки 2 (нижний регистр).
# Для "trigger" поиск идёт с правого конца листа — он в крайнем правом столбце.
_COL_PATTERNS = {
    "date_from":  ["начало период", "дата нач"],
    "date_to":    ["конец период", "дата оконч"],
    "client":     ["клиент", "покупатель", "контрагент"],
    "ds":         ["основание", "дополнит"],
    "comment":    ["комментарий"],
    "category":   ["категория"],
    "notif_num":  ["номер уведомл", "№ уведомл"],
    "turnover":   ["товарооборот"],
    "exclusion":  ["вывод из-под"],
    "returns":    ["обратн"],
    "overdue":    ["просрочен"],
    "samples":    ["образц"],
    "base":       ["база для начисл", "база начисл"],
    "rate":       ["премия, %", "% премии"],
    "bonus":      ["сумма премии"],
    "trigger":    ["уведомление"],
}


def detect_columns(ws_calc) -> dict:
    """
    Читает строку 2 листа «Расчет премии» и ищет столбцы по ключевым словам.
    Возвращает dict {ключ: индекс_0based}.
    Для не найденных столбцов подставляет значения из _COL_DEFAULTS.
    """
    try:
        headers = [(cell.value or "").strip().lower() for cell in ws_calc[2]]
    except Exception:
        return dict(_COL_DEFAULTS)

    cols = {}
    for key, patterns in _COL_PATTERNS.items():
        # Триггер ищем с правого конца — он в последнем столбце
        search_range = (
            range(len(headers) - 1, -1, -1)
            if key == "trigger"
            else range(len(headers))
        )
        found_idx = None
        for i in search_range:
            h = headers[i]
            if any(p in h for p in patterns):
                found_idx = i
                break
        cols[key] = found_idx if found_idx is not None else _COL_DEFAULTS[key]

    return cols


# ---------------------------------------------------------------------------
# Локализация
# ---------------------------------------------------------------------------

MONTHS_RU = {
    1: "января",  2: "февраля", 3: "марта",    4: "апреля",
    5: "мая",     6: "июня",    7: "июля",     8: "августа",
    9: "сентября",10: "октября",11: "ноября",  12: "декабря",
}


# ---------------------------------------------------------------------------
# Форматирование дат
# ---------------------------------------------------------------------------

def fmt_date_title(dt) -> str:
    """25 мая 2026 года — для заголовка уведомления."""
    return f"{dt.day} {MONTHS_RU[dt.month]} {dt.year} года"


def fmt_date_full(dt) -> str:
    """«01» января 2026 — для тела документа."""
    return f"«{dt.day:02d}» {MONTHS_RU[dt.month]} {dt.year}"


# ---------------------------------------------------------------------------
# Разбор ДС
# ---------------------------------------------------------------------------

def parse_ds(ds_str: str):
    """
    'ДС 146 от 13.02.2026' → ('146', datetime(2026,2,13))
    'ДС №105 от 13.02.26'  → ('105', datetime(2026,2,13))
    """
    m = re.search(
        r"ДС\s*№?\s*(\d+)\s+от\s+(\d{1,2})\.(\d{1,2})\.(\d{2,4})",
        str(ds_str)
    )
    if not m:
        return "?", datetime.today()
    num  = m.group(1)
    day, month, year = int(m.group(2)), int(m.group(3)), int(m.group(4))
    if year < 100:
        year += 2000
    return num, datetime(year, month, day)


# ---------------------------------------------------------------------------
# Форматирование чисел
# ---------------------------------------------------------------------------

def safe_float(val) -> float:
    """Пустое / прочерк / None → 0.0, иначе float."""
    if val is None or str(val).strip() in ("", "-", " -"):
        return 0.0
    try:
        return float(val)
    except (ValueError, TypeError):
        return 0.0


def fmt_money(amount: float) -> str:
    """1 234 567.89 → '1 234 567,89' (пробел-разделитель, запятая)."""
    amount = round(amount, 2)
    integer = int(amount)
    kopecks = round((amount - integer) * 100)
    int_str = f"{integer:,}".replace(",", " ")
    return f"{int_str},{kopecks:02d}"


def fmt_cell(val) -> str:
    """0 / пусто → '-', иначе как деньги."""
    f = safe_float(val)
    return "-" if f == 0 else fmt_money(f)


def fmt_percent(rate: float) -> str:
    """0.08 → '8%', 0.0015 → '0,15%', 0.005 → '0,5%'"""
    pct = rate * 100
    if abs(pct - round(pct)) < 1e-9:
        return f"{int(round(pct))}%"
    s = f"{pct:.6f}".rstrip("0").rstrip(".")
    return f"{s.replace('.', ',')}%"


def _ruble_form(n: int) -> str:
    n = abs(n) % 100
    if 11 <= n <= 19:
        return "рублей"
    n = n % 10
    return {1: "рубль", 2: "рубля", 3: "рубля", 4: "рубля"}.get(n, "рублей")


def _kopeck_form(n: int) -> str:
    n = abs(n) % 100
    if 11 <= n <= 19:
        return "копеек"
    n = n % 10
    return {1: "копейка", 2: "копейки", 3: "копейки", 4: "копейки"}.get(n, "копеек")


def amount_in_words(amount: float) -> str:
    """326 000.0 → 'Триста двадцать шесть тысяч рублей 00 копеек'"""
    amount  = round(amount, 2)
    integer = int(amount)
    kopecks = round((amount - integer) * 100)
    words   = num2words(integer, lang="ru")
    words   = words[0].upper() + words[1:]
    return f"{words} {_ruble_form(integer)} {kopecks:02d} {_kopeck_form(kopecks)}"


def normalize_quotes(name: str) -> str:
    """ООО "Русский Свет" → ООО «Русский Свет»"""
    return re.sub(r'"([^"]+)"', r"«\1»", name)


# ---------------------------------------------------------------------------
# Работа с Word-документом
# ---------------------------------------------------------------------------

def _get_text_width(doc) -> int:
    """Ширина текстовой области в твипах (ширина страницы минус поля)."""
    sectPr = doc.element.body.find(qn("w:sectPr"))
    if sectPr is None:
        return 9638  # A4 с полями 2 см
    pgSz  = sectPr.find(qn("w:pgSz"))
    pgMar = sectPr.find(qn("w:pgMar"))
    page_w = int(pgSz.get(qn("w:w"), 11906)) if pgSz is not None else 11906
    if pgMar is not None:
        ml = int(pgMar.get(qn("w:left"),  1134))
        mr = int(pgMar.get(qn("w:right"), 1134))
    else:
        ml = mr = 1134
    return page_w - ml - mr


def _scale_table_to_width(tbl_element, text_width: int):
    """
    Если суммарная ширина столбцов таблицы превышает text_width,
    пропорционально уменьшает все столбцы и ячейки.
    """
    tblGrid = tbl_element.find(qn("w:tblGrid"))
    if tblGrid is None:
        return
    gridCols = tblGrid.findall(qn("w:gridCol"))
    if not gridCols:
        return

    col_widths = [int(col.get(qn("w:w"), 0)) for col in gridCols]
    total_w = sum(col_widths)
    if total_w == 0 or total_w <= text_width:
        return

    scale = text_width / total_w
    print(f"  [таблица] масштаб {total_w}→{text_width} twips (×{scale:.3f})")

    # Масштабируем сетку столбцов
    new_widths = [max(1, round(w * scale)) for w in col_widths]
    for col, new_w in zip(gridCols, new_widths):
        col.set(qn("w:w"), str(new_w))

    # Масштабируем ширины ячеек
    for tr in tbl_element.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None and tcW.get(qn("w:type"), "dxa") == "dxa":
                w = int(tcW.get(qn("w:w"), 0))
                tcW.set(qn("w:w"), str(max(1, round(w * scale))))

    # Устанавливаем ширину таблицы = 100% текстовой области (AutoFit to Window)
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is not None:
            tblW.set(qn("w:w"), "5000")
            tblW.set(qn("w:type"), "pct")
        else:
            el = OxmlElement("w:tblW")
            el.set(qn("w:w"), "5000")
            el.set(qn("w:type"), "pct")
            tblPr.append(el)


def _fix_body_order(doc):
    """
    1. Убрать плавающее позиционирование (w:tblpPr) у таблицы данных,
       сбросить левый отступ (w:tblInd = 0).
    2. Масштабировать столбцы, если таблица шире текстовой области.
    3. Переставить таблицу перед абзацем «Итого» (paras[6]), если нужно.
    """
    body = doc.element.body
    all_children = list(body)

    tbls = [ch for ch in all_children if ch.tag.split("}")[1] == "tbl"]
    if len(tbls) < 2:
        return
    data_tbl = tbls[1]

    tblPr = data_tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        # Убрать плавающую позицию
        tblpPr = tblPr.find(qn("w:tblpPr"))
        if tblpPr is not None:
            tblPr.remove(tblpPr)

        # Сбросить левый отступ до нуля
        tblInd = tblPr.find(qn("w:tblInd"))
        if tblInd is not None:
            tblInd.set(qn("w:w"), "0")
            tblInd.set(qn("w:type"), "dxa")
        else:
            el = OxmlElement("w:tblInd")
            el.set(qn("w:w"), "0")
            el.set(qn("w:type"), "dxa")
            tblPr.append(el)

    # Масштабировать столбцы, если таблица не помещается по ширине
    _scale_table_to_width(data_tbl, _get_text_width(doc))

    # Переставить таблицу перед абзацем «Итого...» (paras[6])
    body_paras = [ch for ch in all_children if ch.tag.split("}")[1] == "p"]
    if len(body_paras) < 7:
        return
    total_p = body_paras[6]

    tbl_idx   = all_children.index(data_tbl)
    total_idx = all_children.index(total_p)

    if tbl_idx > total_idx:
        body.remove(data_tbl)
        body.insert(total_idx, data_tbl)


def _clear_para_content(para):
    """Удалить все раны и маркеры проверки орфографии, оставить pPr."""
    p = para._p
    remove_tags = {
        qn("w:r"),
        qn("w:proofErr"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:del"),
        qn("w:ins"),
    }
    for child in list(p):
        if child.tag in remove_tags:
            p.remove(child)


def _set_para_text(para, text: str):
    """
    Заменить текст параграфа на text, сохранив форматирование
    (шрифт, размер, жирность) первого содержательного рана.
    """
    font_name, font_size, bold = "Arial", None, None
    for run in para.runs:
        if run.font.name or run.bold is not None or run.font.size:
            font_name = run.font.name or font_name
            font_size = run.font.size or font_size
            bold      = run.bold
            break

    _clear_para_content(para)
    run           = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold      = bold


def _set_cell_text(cell, text: str, bold: bool = None):
    """Заменить текст первого параграфа ячейки таблицы."""
    para = cell.paragraphs[0]

    font_name, font_size = "Arial", None
    orig_bold = None
    for run in para.runs:
        if run.font.name or run.font.size or run.bold is not None:
            font_name = run.font.name or font_name
            font_size = run.font.size or font_size
            orig_bold = run.bold
            break

    _clear_para_content(para)
    run           = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold      = orig_bold if bold is None else bold


# ---------------------------------------------------------------------------
# Чтение Excel
# ---------------------------------------------------------------------------

def load_excel(excel_path: str):
    """
    Возвращает:
        rows      — список строк с отметкой «Сделать»
        contracts — {клиент: реквизиты договора}
        legend    — {комментарий: наименование премии}
        cols      — {ключ: индекс_0based} (определяется по заголовкам строки 2)
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    # Реестр договоров
    ws_reg    = wb["Реестр договоров"]
    contracts = {}
    for row in ws_reg.iter_rows(min_row=2, values_only=True):
        if row[0] and row[1]:
            contracts[str(row[0]).strip()] = str(row[1]).strip()

    # Легенда наименований премий
    ws_legend = wb["Наименование премии_легенда"]
    legend    = {}
    for row in ws_legend.iter_rows(min_row=2, values_only=True):
        if row[0] and row[1]:
            legend[str(row[0]).strip()] = str(row[1]).strip()
    if "Неликвид" in legend:
        legend["Неликвиды"] = legend["Неликвид"]

    # Определяем столбцы по заголовкам строки 2
    ws_calc = wb["Расчет премии"]
    cols    = detect_columns(ws_calc)

    # Строки для генерации
    rows = []
    for row in ws_calc.iter_rows(min_row=3, values_only=True):
        trigger_idx = cols["trigger"]
        trigger = row[trigger_idx] if trigger_idx < len(row) else None
        if trigger and str(trigger).strip().lower() == "сделать":
            rows.append(row)

    return rows, contracts, legend, cols


# ---------------------------------------------------------------------------
# Генерация одного документа
# ---------------------------------------------------------------------------

def build_notification(row, contracts: dict, legend: dict,
                        cols: dict, output_dir: Path, today: date) -> tuple:
    """
    Создаёт .docx для одной строки Excel.
    Возвращает (путь_к_файлу, сумма_премии).
    """
    def get(key):
        idx = cols[key]
        return row[idx] if idx < len(row) else None

    notif_num = str(get("notif_num")).strip()
    client    = str(get("client")).strip()
    ds_str    = str(get("ds")).strip() if get("ds") else ""
    comment   = str(get("comment")).strip() if get("comment") else ""
    category  = str(get("category")).strip() if get("category") else ""
    date_from = get("date_from")
    date_to   = get("date_to")

    # Дата уведомления
    is_rs       = "русский свет" in client.lower()
    notif_date  = date_to if is_rs else datetime(today.year, today.month, today.day)

    # ДС
    ds_num, ds_date = parse_ds(ds_str)

    # Договор
    contract = contracts.get(client)
    if contract is None:
        contract = contracts.get(normalize_quotes(client), "— нет в реестре —")

    # Наименование премии
    bonus_name = legend.get(comment) or legend.get(category, comment)

    # Маркетинг-фраза
    is_marketing   = "маркетинг" in category.lower()
    marketing_frag = "за достигнутый товарооборот " if is_marketing else ""

    # Суммы берём напрямую из таблицы Excel (кэшированные значения формул)
    turnover  = safe_float(get("turnover"))   # I
    exclusion = safe_float(get("exclusion"))  # II
    returns_  = safe_float(get("returns"))    # III
    overdue   = safe_float(get("overdue"))    # IV
    base      = safe_float(get("base"))       # V  — формула в Excel (R-S-T-U)
    rate      = safe_float(get("rate"))       # VI
    samples   = safe_float(get("samples"))    # VII
    bonus     = safe_float(get("bonus"))      # VIII — формула в Excel (V*VI-VII)

    # Текст параграфов
    client_display = normalize_quotes(client)

    title_text = (
        f"Уведомление о расчете премии {notif_num} "
        f"от {fmt_date_title(notif_date)}"
    )

    body_text = (
        f"        Настоящим уведомляем, что в соответствии с Дополнительным "
        f"соглашением № {ds_num} от {fmt_date_full(ds_date)} г. "
        f"к Договору {contract} "
        f"между АО «ЛЕДВАНС» и {client_display} "
        f"за период с {fmt_date_full(date_from)} года "
        f"по {fmt_date_full(date_to)} года "
        f"{marketing_frag}начислена премия:"
    )

    total_text = (
        f"          Итого размер премии составил {fmt_money(bonus)} руб. "
        f"({amount_in_words(bonus)}) без НДС. Премия НДС не облагается. "
    )

    # Копируем шаблон и заменяем содержимое
    doc   = Document(str(TEMPLATE_PATH))
    _fix_body_order(doc)
    paras = doc.paragraphs

    _set_para_text(paras[2], title_text)
    _set_para_text(paras[4], body_text)
    _set_para_text(paras[6], total_text)

    # Таблица данных (doc.tables[1])
    t        = doc.tables[1]
    data_row = t.rows[2]
    itg_row  = t.rows[3]

    _set_cell_text(data_row.cells[0], bonus_name)
    _set_cell_text(data_row.cells[1], fmt_money(turnover))
    _set_cell_text(data_row.cells[2], fmt_cell(exclusion))
    _set_cell_text(data_row.cells[3], fmt_cell(returns_))
    _set_cell_text(data_row.cells[4], fmt_cell(overdue))
    _set_cell_text(data_row.cells[5], fmt_money(base))
    _set_cell_text(data_row.cells[6], fmt_percent(rate))
    _set_cell_text(data_row.cells[7], fmt_cell(samples))
    _set_cell_text(data_row.cells[8], fmt_money(bonus))

    _set_cell_text(itg_row.cells[8], fmt_money(bonus), bold=True)

    # Сохраняем — имя файла совпадает с заголовком документа
    notif_clean = re.sub(r'[/\s]+', '', notif_num)
    file_name   = (
        f"Уведомление о расчете премии {notif_clean} "
        f"от {notif_date.day} {MONTHS_RU[notif_date.month]} {notif_date.year} года.docx"
    )
    safe_file   = re.sub(r'[\\:*?"<>|]', '_', file_name)
    out_path    = output_dir / safe_file
    doc.save(str(out_path))
    return out_path, bonus


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    excel_path = sys.argv[1]
    if not os.path.isfile(excel_path):
        sys.exit(f"Ошибка: файл не найден — {excel_path}")

    if not TEMPLATE_PATH.is_file():
        sys.exit(f"Ошибка: шаблон не найден — {TEMPLATE_PATH}\n"
                 f"Положите template.docx рядом со скриптом.")

    print(f"generate.py версия {VERSION}")
    today                         = date.today()
    rows, contracts, legend, cols = load_excel(excel_path)

    if not rows:
        print("Строк с отметкой «Сделать» в столбце «Уведомление» не найдено.")
        sys.exit(0)

    output_dir = Path(excel_path).parent / f"Уведомления_{today.strftime('%Y-%m-%d')}"
    output_dir.mkdir(exist_ok=True)

    print(f"Найдено строк: {len(rows)}")
    print(f"Папка вывода:  {output_dir}\n")

    ok = err = 0
    for row in rows:
        notif_num = str(row[cols["notif_num"]]).strip() if row[cols["notif_num"]] else "?"
        client    = str(row[cols["client"]]).strip()    if row[cols["client"]]    else "?"
        try:
            out_path, bonus = build_notification(
                row, contracts, legend, cols, output_dir, today
            )
            print(f"  ✓  {notif_num:<30}  {client:<40}  {fmt_money(bonus)} руб.")
            ok += 1
        except Exception as exc:
            print(f"  ✗  {notif_num:<30}  ОШИБКА: {exc}")
            err += 1

    print(f"\nГотово: {ok} создано, {err} ошибок.")
    if ok:
        print(f"Документы: {output_dir}")


if __name__ == "__main__":
    main()
