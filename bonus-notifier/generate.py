#!/usr/bin/env python3
"""
Генератор уведомлений о расчёте премии.

Использование:
    python generate.py путь_к_excel.xlsx

В Excel на листе «Расчет премии» укажите «Сделать» в столбце «Уведомление»
для строк, по которым нужно создать документ.

Файл template.docx должен лежать рядом со скриптом.
"""

import sys
import re
import os
from datetime import date, datetime
from pathlib import Path
import warnings

warnings.filterwarnings("ignore", category=UserWarning)

try:
    import openpyxl
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")

try:
    from num2words import num2words
except ImportError:
    sys.exit("Установите зависимости: pip install openpyxl python-docx num2words")


# ---------------------------------------------------------------------------
# Пути
# ---------------------------------------------------------------------------

SCRIPT_DIR    = Path(__file__).parent
TEMPLATE_PATH = SCRIPT_DIR / "template.docx"


# ---------------------------------------------------------------------------
# Индексы столбцов листа «Расчет премии» (0-based, используются как fallback)
# Скрипт ищет столбцы «Уведомление», «База» и «Сумма премии» по заголовку,
# поэтому добавление новых столбцов в таблицу не влияет на работу скрипта.
# ---------------------------------------------------------------------------

COL_DATE_FROM = 2   # C  — начало периода
COL_DATE_TO   = 3   # D  — конец периода
COL_CLIENT    = 5   # F  — клиент
COL_DS        = 6   # G  — основание (ДС №... от ...)
COL_COMMENT   = 7   # H  — комментарий (ключ легенды)
COL_CATEGORY  = 8   # I  — категория (для маркетинг-фразы)
COL_NOTIF_NUM = 15  # P  — номер уведомления
COL_TURNOVER  = 17  # R  — товарооборот I
COL_EXCLUSION = 18  # S  — вывод из-под бонуса II
COL_RETURNS   = 19  # T  — обратная реализация III
COL_OVERDUE   = 20  # U  — просроченная задолженность IV
COL_SAMPLES   = 21  # V  — отгруженные образцы VII
COL_RATE      = 23  # X  — процент премии VI
COL_BASE_DEF  = 22  # W  — база для начисления V  (fallback)
COL_BONUS_DEF = 24  # Y  — сумма премии VIII       (fallback)
COL_TRIGGER   = 25  # Z  — «Уведомление» (триггер, fallback)


# ---------------------------------------------------------------------------
# Локализация
# ---------------------------------------------------------------------------

MONTHS_RU = {
    1: "января",  2: "февраля", 3: "марта",    4: "апреля",
    5: "мая",     6: "июня",    7: "июля",     8: "августа",
    9: "сентября",10: "октября",11: "ноября",  12: "декабря",
}


# ---------------------------------------------------------------------------
# Форматирование дат
# ---------------------------------------------------------------------------

def fmt_date_title(dt) -> str:
    """25 мая 2026 года — для заголовка уведомления."""
    return f"{dt.day} {MONTHS_RU[dt.month]} {dt.year} года"


def fmt_date_full(dt) -> str:
    """«01» января 2026 — для тела документа."""
    return f"«{dt.day:02d}» {MONTHS_RU[dt.month]} {dt.year}"


# ---------------------------------------------------------------------------
# Разбор ДС
# ---------------------------------------------------------------------------

def parse_ds(ds_str: str):
    """
    'ДС 146 от 13.02.2026' → ('146', datetime(2026,2,13))
    'ДС №105 от 13.02.26'  → ('105', datetime(2026,2,13))
    """
    m = re.search(
        r"ДС\s*№?\s*(\d+)\s+от\s+(\d{1,2})\.(\d{1,2})\.(\d{2,4})",
        str(ds_str)
    )
    if not m:
        return "?", datetime.today()
    num  = m.group(1)
    day, month, year = int(m.group(2)), int(m.group(3)), int(m.group(4))
    if year < 100:
        year += 2000
    return num, datetime(year, month, day)


# ---------------------------------------------------------------------------
# Форматирование чисел
# ---------------------------------------------------------------------------

def safe_float(val) -> float:
    """Пустое / прочерк / None → 0.0, иначе float."""
    if val is None or str(val).strip() in ("", "-", " -"):
        return 0.0
    try:
        return float(val)
    except (ValueError, TypeError):
        return 0.0


def fmt_money(amount: float) -> str:
    """1 234 567.89 → '1 234 567,89' (пробел-разделитель, запятая)."""
    amount = round(amount, 2)
    integer = int(amount)
    kopecks = round((amount - integer) * 100)
    int_str = f"{integer:,}".replace(",", " ")
    return f"{int_str},{kopecks:02d}"


def fmt_cell(val) -> str:
    """0 / пусто → '-', иначе как деньги."""
    f = safe_float(val)
    return "-" if f == 0 else fmt_money(f)


def fmt_percent(rate: float) -> str:
    """0.08 → '8%', 0.0015 → '0,15%', 0.005 → '0,5%'"""
    pct = rate * 100
    if abs(pct - round(pct)) < 1e-9:
        return f"{int(round(pct))}%"
    s = f"{pct:.6f}".rstrip("0").rstrip(".")
    return f"{s.replace('.', ',')}%"


def _ruble_form(n: int) -> str:
    n = abs(n) % 100
    if 11 <= n <= 19:
        return "рублей"
    n = n % 10
    return {1: "рубль", 2: "рубля", 3: "рубля", 4: "рубля"}.get(n, "рублей")


def _kopeck_form(n: int) -> str:
    n = abs(n) % 100
    if 11 <= n <= 19:
        return "копеек"
    n = n % 10
    return {1: "копейка", 2: "копейки", 3: "копейки", 4: "копейки"}.get(n, "копеек")


def amount_in_words(amount: float) -> str:
    """326 000.0 → 'Триста двадцать шесть тысяч рублей 00 копеек'"""
    amount  = round(amount, 2)
    integer = int(amount)
    kopecks = round((amount - integer) * 100)
    words   = num2words(integer, lang="ru")
    words   = words[0].upper() + words[1:]
    return f"{words} {_ruble_form(integer)} {kopecks:02d} {_kopeck_form(kopecks)}"


def normalize_quotes(name: str) -> str:
    """ООО "Русский Свет" → ООО «Русский Свет»"""
    return re.sub(r'"([^"]+)"', r"«\1»", name)


# ---------------------------------------------------------------------------
# Поиск столбцов по заголовку
# ---------------------------------------------------------------------------

def _find_column(ws, keywords, default, max_header_row=3):
    """Вернуть 0-based индекс столбца, заголовок которого содержит все ключевые слова.

    Просматривает строки 1..max_header_row. Если не найдено — возвращает default.
    Благодаря этому добавление новых столбцов в таблицу не ломает скрипт.
    """
    for row in ws.iter_rows(min_row=1, max_row=max_header_row, values_only=True):
        for col_i, val in enumerate(row):
            if val is None:
                continue
            text = str(val).lower()
            if all(kw.lower() in text for kw in keywords):
                return col_i
    return default


# ---------------------------------------------------------------------------
# Работа с Word-документом
# ---------------------------------------------------------------------------

def _fix_body_order(doc):
    """Убрать плавающее позиционирование, поставить таблицу по центру страницы
    через отрицательный отступ, и при необходимости переставить её перед «Итого».

    Исходная таблица шире текстовой области (11319 vs 9072 twips) и была
    плавающей. Вместо масштабирования — центрируем на странице через tblInd,
    сохраняя оригинальные ширины колонок.
    """
    body = doc.element.body
    all_children = list(body)

    # Второй <w:tbl> = таблица данных
    tbls = [ch for ch in all_children if ch.tag.split("}")[1] == "tbl"]
    if len(tbls) < 2:
        return
    data_tbl = tbls[1]

    tblPr = data_tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        # 1. Убрать плавающее позиционирование
        tblpPr = tblPr.find(qn("w:tblpPr"))
        if tblpPr is not None:
            tblPr.remove(tblpPr)

        # 2. Считать размеры страницы
        sectPr  = body.find(qn("w:sectPr"))
        page_w  = 11906
        left_m  = 1418
        right_m = 1416
        if sectPr is not None:
            pgSz  = sectPr.find(qn("w:pgSz"))
            pgMar = sectPr.find(qn("w:pgMar"))
            if pgSz  is not None:
                page_w  = int(pgSz.get(qn("w:w"),    page_w))
            if pgMar is not None:
                left_m  = int(pgMar.get(qn("w:left"),  left_m))
                right_m = int(pgMar.get(qn("w:right"), right_m))

        # 3. Ширина таблицы
        tblW      = tblPr.find(qn("w:tblW"))
        tbl_width = int(tblW.get(qn("w:w"), 0))     if tblW is not None else 0
        cur_type  = tblW.get(qn("w:type"), "dxa")   if tblW is not None else "dxa"
        text_w    = page_w - left_m - right_m        # 9072

        if cur_type == "dxa" and tbl_width > text_w:
            # Таблица шире текстовой области — центрируем через отрицательный tblInd
            # Левый край от края страницы: (page_w - tbl_width) / 2
            left_edge = (page_w - tbl_width) // 2    # 293 twips от края страницы
            tbl_ind   = left_edge - left_m            # отступ от начала текстовой области (< 0)

            old_ind = tblPr.find(qn("w:tblInd"))
            if old_ind is not None:
                tblPr.remove(old_ind)
            new_ind = OxmlElement("w:tblInd")
            new_ind.set(qn("w:w"),   str(tbl_ind))
            new_ind.set(qn("w:type"), "dxa")
            # Вставить сразу после tblW
            tblPr.insert(list(tblPr).index(tblW) + 1, new_ind)

    # 4. Переставить таблицу перед абзацем «Итого» (paras[6]), если нужно
    body_paras = [ch for ch in all_children if ch.tag.split("}")[1] == "p"]
    if len(body_paras) < 7:
        return
    total_p = body_paras[6]

    tbl_idx   = all_children.index(data_tbl)
    total_idx = all_children.index(total_p)

    if tbl_idx > total_idx:
        body.remove(data_tbl)
        body.insert(total_idx, data_tbl)


def _clear_para_content(para):
    """Удалить все раны и маркеры проверки орфографии, оставить pPr."""
    p = para._p
    remove_tags = {
        qn("w:r"),
        qn("w:proofErr"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:del"),
        qn("w:ins"),
    }
    for child in list(p):
        if child.tag in remove_tags:
            p.remove(child)


def _set_para_text(para, text: str):
    """Заменить текст параграфа, сохранив форматирование первого рана."""
    font_name, font_size, bold = "Arial", None, None
    for run in para.runs:
        if run.font.name or run.bold is not None or run.font.size:
            font_name = run.font.name or font_name
            font_size = run.font.size or font_size
            bold      = run.bold
            break

    _clear_para_content(para)
    run           = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold      = bold


def _set_cell_text(cell, text: str, bold: bool = None):
    """Заменить текст первого параграфа ячейки таблицы."""
    para = cell.paragraphs[0]

    font_name, font_size = "Arial", None
    orig_bold = None
    for run in para.runs:
        if run.font.name or run.font.size or run.bold is not None:
            font_name = run.font.name or font_name
            font_size = run.font.size or font_size
            orig_bold = run.bold
            break

    _clear_para_content(para)
    run           = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold      = orig_bold if bold is None else bold


# ---------------------------------------------------------------------------
# Чтение Excel
# ---------------------------------------------------------------------------

def load_excel(excel_path: str):
    """
    Возвращает:
        rows      — список строк с отметкой «Сделать»
        contracts — {клиент: реквизиты договора}
        legend    — {комментарий: наименование премии}
        base_col  — индекс столбца «База для начисления»
        bonus_col — индекс столбца «Сумма премии»
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    # Реестр договоров
    ws_reg    = wb["Реестр договоров"]
    contracts = {}
    for row in ws_reg.iter_rows(min_row=2, values_only=True):
        if row[0] and row[1]:
            contracts[str(row[0]).strip()] = str(row[1]).strip()

    # Легенда наименований премий
    ws_legend = wb["Наименование премии_легенда"]
    legend    = {}
    for row in ws_legend.iter_rows(min_row=2, values_only=True):
        if row[0] and row[1]:
            legend[str(row[0]).strip()] = str(row[1]).strip()
    if "Неликвид" in legend:
        legend["Неликвиды"] = legend["Неликвид"]

    # Найти ключевые столбцы по заголовку (устойчиво к добавлению новых столбцов)
    ws_calc   = wb["Расчет премии"]
    trig_col  = _find_column(ws_calc, ["уведомление"],       COL_TRIGGER)
    base_col  = _find_column(ws_calc, ["база", "начисл"],    COL_BASE_DEF)
    bonus_col = _find_column(ws_calc, ["сумма", "премии"],   COL_BONUS_DEF)

    rows = []
    for row in ws_calc.iter_rows(min_row=3, values_only=True):
        trig = row[trig_col] if trig_col < len(row) else None
        if trig and str(trig).strip().lower() == "сделать":
            rows.append(row)

    return rows, contracts, legend, base_col, bonus_col


# ---------------------------------------------------------------------------
# Генерация одного документа
# ---------------------------------------------------------------------------

def build_notification(row, contracts: dict, legend: dict,
                        output_dir: Path, today: date,
                        base_col: int, bonus_col: int) -> tuple:
    """
    Создаёт .docx для одной строки Excel.
    Возвращает (путь_к_файлу, сумма_премии).
    """
    notif_num = str(row[COL_NOTIF_NUM]).strip()
    client    = str(row[COL_CLIENT]).strip()
    ds_str    = str(row[COL_DS]).strip() if row[COL_DS] else ""
    comment   = str(row[COL_COMMENT]).strip() if row[COL_COMMENT] else ""
    category  = str(row[COL_CATEGORY]).strip() if row[COL_CATEGORY] else ""
    date_from = row[COL_DATE_FROM]
    date_to   = row[COL_DATE_TO]

    # Дата уведомления
    is_rs      = "русский свет" in client.lower()
    notif_date = date_to if is_rs else datetime(today.year, today.month, today.day)

    # ДС
    ds_num, ds_date = parse_ds(ds_str)

    # Договор
    contract = contracts.get(client)
    if contract is None:
        contract = contracts.get(normalize_quotes(client), "— нет в реестре —")

    # Наименование премии
    bonus_name = legend.get(comment) or legend.get(category, comment)

    # Маркетинг-фраза
    is_marketing   = "маркетинг" in category.lower()
    marketing_frag = "за достигнутый товарооборот " if is_marketing else ""

    # Значения из Excel (кэшированные формулы)
    turnover  = safe_float(row[COL_TURNOVER])   # I
    exclusion = safe_float(row[COL_EXCLUSION])  # II
    returns_  = safe_float(row[COL_RETURNS])    # III
    overdue   = safe_float(row[COL_OVERDUE])    # IV
    samples   = safe_float(row[COL_SAMPLES])    # VII
    rate      = safe_float(row[COL_RATE])       # VI
    base      = safe_float(row[base_col]  if base_col  < len(row) else None)   # V
    bonus     = round(safe_float(row[bonus_col] if bonus_col < len(row) else None), 2)  # VIII

    # Текст параграфов
    client_display = normalize_quotes(client)

    title_text = (
        f"Уведомление о расчете премии {notif_num} "
        f"от {fmt_date_title(notif_date)}"
    )

    body_text = (
        f"        Настоящим уведомляем, что в соответствии с Дополнительным "
        f"соглашением № {ds_num} от {fmt_date_full(ds_date)} г. "
        f"к Договору {contract} "
        f"между АО «ЛЕДВАНС» и {client_display} "
        f"за период с {fmt_date_full(date_from)} года "
        f"по {fmt_date_full(date_to)} года "
        f"{marketing_frag}начислена премия:"
    )

    total_text = (
        f"          Итого размер премии составил {fmt_money(bonus)} руб. "
        f"({amount_in_words(bonus)}) без НДС. Премия НДС не облагается. "
    )

    # Копируем шаблон и заменяем содержимое
    doc   = Document(str(TEMPLATE_PATH))
    _fix_body_order(doc)    # убирает плавающую таблицу и выставляет ширину 100%
    paras = doc.paragraphs

    _set_para_text(paras[2], title_text)
    _set_para_text(paras[4], body_text)
    _set_para_text(paras[6], total_text)

    # Таблица данных (doc.tables[1])
    t        = doc.tables[1]
    data_row = t.rows[2]
    itg_row  = t.rows[3]

    _set_cell_text(data_row.cells[0], bonus_name)
    _set_cell_text(data_row.cells[1], fmt_money(turnover))
    _set_cell_text(data_row.cells[2], fmt_cell(exclusion))
    _set_cell_text(data_row.cells[3], fmt_cell(returns_))
    _set_cell_text(data_row.cells[4], fmt_cell(overdue))
    _set_cell_text(data_row.cells[5], fmt_money(base))
    _set_cell_text(data_row.cells[6], fmt_percent(rate))
    _set_cell_text(data_row.cells[7], fmt_cell(samples))
    _set_cell_text(data_row.cells[8], fmt_money(bonus))

    _set_cell_text(itg_row.cells[8], fmt_money(bonus), bold=True)

    notif_clean = re.sub(r'[/\s]+', '', notif_num)
    file_name   = (
        f"Уведомление о расчете премии {notif_clean} "
        f"от {notif_date.day} {MONTHS_RU[notif_date.month]} {notif_date.year} года.docx"
    )
    safe_file = re.sub(r'[\\:*?"<>|]', '_', file_name)
    out_path  = output_dir / safe_file
    doc.save(str(out_path))
    return out_path, bonus


# ---------------------------------------------------------------------------
# Точка входа
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    excel_path = sys.argv[1]
    if not os.path.isfile(excel_path):
        sys.exit(f"Ошибка: файл не найден — {excel_path}")

    if not TEMPLATE_PATH.is_file():
        sys.exit(f"Ошибка: шаблон не найден — {TEMPLATE_PATH}\n"
                 f"Положите template.docx рядом со скриптом.")

    today = date.today()
    rows, contracts, legend, base_col, bonus_col = load_excel(excel_path)

    if not rows:
        print("Строк с отметкой «Сделать» в столбце «Уведомление» не найдено.")
        sys.exit(0)

    output_dir = Path(excel_path).parent / f"Уведомления_{today.strftime('%Y-%m-%d')}"
    output_dir.mkdir(exist_ok=True)

    print(f"Найдено строк: {len(rows)}")
    print(f"Папка вывода:  {output_dir}\n")

    ok = err = 0
    for row in rows:
        notif_num = str(row[COL_NOTIF_NUM]).strip() if row[COL_NOTIF_NUM] else "?"
        client    = str(row[COL_CLIENT]).strip()    if row[COL_CLIENT]    else "?"
        try:
            out_path, bonus = build_notification(
                row, contracts, legend, output_dir, today, base_col, bonus_col
            )
            print(f"  ✓  {notif_num:<30}  {client:<40}  {fmt_money(bonus)} руб.")
            ok += 1
        except Exception as exc:
            print(f"  ✗  {notif_num:<30}  ОШИБКА: {exc}")
            err += 1

    print(f"\nГотово: {ok} создано, {err} ошибок.")
    if ok:
        print(f"Документы: {output_dir}")


if __name__ == "__main__":
    main()
